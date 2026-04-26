"""Programmatically generate the 3 sample input docs (intentionally inconsistent) and master.docx (canonical styles).

Run once after `uv sync` to bootstrap:

    uv run python scripts/generate_samples.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

SUBDIR = Path(__file__).resolve().parents[1]
INPUT_DIR = SUBDIR / "input"
MASTER_PATH = SUBDIR / "master.docx"


def _make_master() -> None:
    doc = Document()
    style_specs = [
        ("Title", 28, True, RGBColor(0x1F, 0x4E, 0x79)),
        ("Heading 1", 18, True, RGBColor(0x2E, 0x74, 0xB5)),
        ("Heading 2", 14, True, RGBColor(0x44, 0x44, 0x44)),
        ("Normal", 11, False, RGBColor(0x00, 0x00, 0x00)),
    ]
    for name, font_size, bold, color in style_specs:
        style = doc.styles[name]
        style.font.size = Pt(font_size)
        style.font.bold = bold
        style.font.color.rgb = color
    doc.add_paragraph("Master Template", style="Title")
    doc.add_paragraph("Section", style="Heading 1")
    doc.add_paragraph("Subsection", style="Heading 2")
    doc.add_paragraph("Body text in Normal style.", style="Normal")
    doc.add_table(rows=1, cols=2, style="Table Grid")
    doc.save(str(MASTER_PATH))


def _make_sop() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Standard Operating Procedure: New Employee Onboarding")
    r.bold = True
    r.font.size = Pt(20)
    doc.add_paragraph("This SOP defines onboarding for new hires.")
    doc.add_paragraph("PURPOSE")
    doc.add_paragraph("Establish a consistent onboarding process across teams.")
    doc.add_paragraph("Steps:")
    doc.add_paragraph("1. HR sends welcome email")
    doc.add_paragraph("2. IT provisions accounts")
    doc.add_paragraph("3. Manager schedules orientation")
    doc.add_paragraph("Owner: People Ops. Effective: 2026-01-01.")
    doc.save(str(INPUT_DIR / "sample_sop.docx"))


def _make_policy() -> None:
    doc = Document()
    doc.add_heading("Acceptable Use Policy", level=0)
    doc.add_paragraph("Scope: All employees and contractors.", style="Normal")
    doc.add_paragraph("Purpose")
    doc.add_paragraph("Define what is and is not allowed on company systems.")
    doc.add_paragraph("Definitions:")
    doc.add_paragraph("- Personal Use: any non-business use of company resources.")
    doc.add_paragraph("- Confidential Data: data marked as such by Data Owner.")
    doc.add_paragraph("Responsibilities: IT enforces, Legal interprets.")
    doc.add_paragraph("Records: signed acknowledgment form, annual training completion log.")
    doc.add_paragraph("References: ISO 27001 Annex A.6, NIST SP 800-53 AC-1.")
    doc.save(str(INPUT_DIR / "sample_policy.docx"))


def _make_procedure() -> None:
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("INCIDENT RESPONSE PROCEDURE")
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph("Owner: SecOps. Approval: CISO.")
    doc.add_paragraph("Steps")
    table = doc.add_table(rows=4, cols=2)
    rows = [
        ("Step", "Action"),
        ("1", "Detect anomaly via SIEM alert"),
        ("2", "Triage and classify severity"),
        ("3", "Engage incident commander"),
    ]
    for i, (a, b) in enumerate(rows):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
    doc.add_paragraph("Records: incident ticket, post-mortem doc.")
    doc.save(str(INPUT_DIR / "sample_procedure.docx"))


def main() -> None:
    INPUT_DIR.mkdir(parents=True, exist_ok=True)
    _make_master()
    _make_sop()
    _make_policy()
    _make_procedure()
    print(f"Wrote master to {MASTER_PATH}")
    print(f"Wrote 3 samples to {INPUT_DIR}/")


if __name__ == "__main__":
    main()
