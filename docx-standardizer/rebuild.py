"""Rebuild a fresh .docx by copying styles from a master template and writing the normalized content section by section."""

from copy import deepcopy
from pathlib import Path
from typing import List

from docx import Document
from docx.document import Document as DocumentObject

from schema import StandardizedDocument

CORE_STYLES = ["Title", "Heading 1", "Heading 2", "Normal", "Table Grid"]


def copy_styles(
    master_path: Path,
    target: DocumentObject,
    names: List[str] = CORE_STYLES,
) -> None:
    """Copy paragraph and table styles from master into target's styles collection.

    Pattern from python-docx maintainer @scanny on Stack Overflow:
    https://stackoverflow.com/questions/49512204
    """
    master = Document(str(master_path))
    src_styles = master.styles
    tgt_styles = target.styles
    src_names = [s.name for s in src_styles]
    tgt_names = [s.name for s in tgt_styles]
    for name in names:
        if name not in src_names:
            continue
        src_style = src_styles[name]
        if name in tgt_names:
            tgt_styles[name].delete()
        tgt_styles.element.append(deepcopy(src_style.element))


def rebuild(
    normalized: StandardizedDocument,
    master_path: Path,
) -> DocumentObject:
    doc = Document()
    copy_styles(master_path, doc)

    doc.add_paragraph(normalized.title, style="Title")

    doc.add_paragraph("Document Control", style="Heading 1")
    dc = normalized.document_control
    dc_table = doc.add_table(rows=4, cols=2, style="Table Grid")
    pairs = [
        ("Doc ID", dc.doc_id),
        ("Owner", dc.owner),
        ("Approval", dc.approval),
        ("Effective Date", dc.effective_date),
    ]
    for i, (k, v) in enumerate(pairs):
        dc_table.rows[i].cells[0].text = k
        dc_table.rows[i].cells[1].text = v

    doc.add_paragraph("Purpose", style="Heading 1")
    doc.add_paragraph(normalized.purpose, style="Normal")

    doc.add_paragraph("Scope", style="Heading 1")
    doc.add_paragraph(normalized.scope, style="Normal")

    if normalized.definitions:
        doc.add_paragraph("Definitions", style="Heading 1")
        for d in normalized.definitions:
            doc.add_paragraph(d.term, style="Heading 2")
            doc.add_paragraph(d.meaning, style="Normal")

    if normalized.responsibilities:
        doc.add_paragraph("Responsibilities", style="Heading 1")
        for r in normalized.responsibilities:
            doc.add_paragraph(r.role, style="Heading 2")
            doc.add_paragraph(r.duties, style="Normal")

    doc.add_paragraph("Procedure", style="Heading 1")
    for s in sorted(normalized.procedure, key=lambda x: x.step_number):
        doc.add_paragraph(f"{s.step_number}. {s.action}", style="Normal")

    if normalized.records:
        doc.add_paragraph("Records", style="Heading 1")
        for r in normalized.records:
            doc.add_paragraph(r, style="Normal")

    if normalized.references:
        doc.add_paragraph("References", style="Heading 1")
        for r in normalized.references:
            doc.add_paragraph(r, style="Normal")

    doc.add_paragraph("Revision History", style="Heading 1")
    rh = doc.add_table(
        rows=len(normalized.revision_history) + 1,
        cols=4,
        style="Table Grid",
    )
    headers = ["Version", "Date", "Author", "Summary"]
    for i, h in enumerate(headers):
        rh.rows[0].cells[i].text = h
    for i, entry in enumerate(normalized.revision_history, start=1):
        rh.rows[i].cells[0].text = entry.version
        rh.rows[i].cells[1].text = entry.date
        rh.rows[i].cells[2].text = entry.author
        rh.rows[i].cells[3].text = entry.summary

    return doc
