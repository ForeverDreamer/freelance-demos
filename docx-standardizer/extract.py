"""Extract title, paragraphs (with style names, bold flag, max font size), tables, and metadata from a .docx file using python-docx. Output is a dict suitable for serialization to OpenAI."""

from pathlib import Path
from typing import Any, Dict, List

from docx import Document


def _paragraph_features(p) -> Dict[str, Any]:
    bolds = [r.bold for r in p.runs if r.bold is not None]
    sizes = [r.font.size.pt for r in p.runs if r.font.size is not None]
    return {
        "bold": any(bolds),
        "max_font_pt": max(sizes) if sizes else None,
    }


def _pick_title(paragraphs: List[Dict[str, Any]], metadata_title: str) -> str:
    """Heading-style first; else largest bold paragraph; else first paragraph."""
    if not paragraphs:
        return metadata_title
    for p in paragraphs:
        if p["style"] in ("Title", "Heading 1"):
            return p["text"]
    bold_with_size = [
        p for p in paragraphs if p["bold"] and p["max_font_pt"] is not None
    ]
    if bold_with_size:
        return max(bold_with_size, key=lambda p: p["max_font_pt"])["text"]
    return paragraphs[0]["text"]


def extract(path: Path) -> Dict[str, Any]:
    doc = Document(str(path))
    paragraphs: List[Dict[str, Any]] = []
    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        paragraphs.append(
            {
                "style": p.style.name if p.style else "Normal",
                "text": p.text,
                **_paragraph_features(p),
            }
        )
    tables: List[List[List[str]]] = []
    for t in doc.tables:
        rows: List[List[str]] = []
        for row in t.rows:
            rows.append([cell.text for cell in row.cells])
        tables.append(rows)
    cp = doc.core_properties
    metadata = {
        "author": cp.author or "",
        "created": cp.created.isoformat() if cp.created else "",
        "modified": cp.modified.isoformat() if cp.modified else "",
        "title": cp.title or "",
    }
    return {
        "title": _pick_title(paragraphs, metadata["title"]),
        "paragraphs": paragraphs,
        "tables": tables,
        "metadata": metadata,
    }
