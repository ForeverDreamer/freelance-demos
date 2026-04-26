import json
from pathlib import Path
from typing import List
from unittest.mock import MagicMock

import pytest
from docx import Document

import normalize as normalize_module
from standardize import process_file


@pytest.fixture
def fixture_input_docx(tmp_path: Path) -> Path:
    p = tmp_path / "in" / "messy.docx"
    p.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph("Some Random Title", style="Title")
    doc.add_paragraph("Body text for testing.", style="Normal")
    doc.save(str(p))
    return p


@pytest.fixture
def fixture_master(tmp_path: Path) -> Path:
    p = tmp_path / "master.docx"
    doc = Document()
    doc.add_paragraph("master ref", style="Title")
    doc.save(str(p))
    return p


def test_pipeline_end_to_end_ok(
    fixture_input_docx: Path,
    fixture_master: Path,
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    output_dir = tmp_path / "out"
    output_dir.mkdir()

    payload = {
        "title": "Normalized Title",
        "document_control": {
            "doc_id": "X",
            "owner": "Y",
            "approval": "Z",
            "effective_date": "2026-01-01",
        },
        "purpose": "Purpose.",
        "scope": "Scope.",
        "definitions": [{"term": "T", "meaning": "M"}],
        "responsibilities": [{"role": "R", "duties": "D"}],
        "procedure": [{"step_number": 1, "action": "Do."}],
        "records": ["rec"],
        "references": ["ref"],
        "revision_history": [
            {
                "version": "1.0",
                "date": "2026-01-01",
                "author": "A",
                "summary": "Init.",
            }
        ],
    }
    mock_client = MagicMock()
    response = MagicMock()
    response.choices = [MagicMock(message=MagicMock(content=json.dumps(payload)))]
    mock_client.chat.completions.create.return_value = response

    monkeypatch.setattr(
        normalize_module, "_build_openai_client", lambda cfg: mock_client
    )

    log_lines: List[str] = []
    status = process_file(
        fixture_input_docx,
        output_dir,
        fixture_master,
        log_lines,
    )

    assert status == "OK"
    assert (output_dir / "messy_standardized.docx").exists()
    assert any("OK" in line for line in log_lines)
