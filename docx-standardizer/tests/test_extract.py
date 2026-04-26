from pathlib import Path

import pytest
from docx import Document
from docx.shared import Pt

from extract import extract


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    p = tmp_path / "sample.docx"
    doc = Document()
    doc.core_properties.author = "Test Author"
    doc.core_properties.title = "Test Title"
    doc.add_paragraph("My Title", style="Title")
    doc.add_paragraph("First Section", style="Heading 1")
    doc.add_paragraph("Body text here.", style="Normal")
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "k1"
    table.rows[0].cells[1].text = "v1"
    table.rows[1].cells[0].text = "k2"
    table.rows[1].cells[1].text = "v2"
    doc.save(str(p))
    return p


@pytest.fixture
def normal_only_docx(tmp_path: Path) -> Path:
    """Mimics sample_sop.docx: title is a hand-rolled bold+large run on a Normal paragraph, every other paragraph is plain Normal."""
    p = tmp_path / "normal_only.docx"
    doc = Document()
    title_para = doc.add_paragraph()
    run = title_para.add_run("Standard Operating Procedure: Onboarding")
    run.bold = True
    run.font.size = Pt(20)
    doc.add_paragraph("This SOP defines onboarding for new hires.")
    doc.add_paragraph("PURPOSE")
    doc.add_paragraph("Establish a consistent onboarding process.")
    doc.save(str(p))
    return p


def test_extract_returns_title(sample_docx: Path) -> None:
    out = extract(sample_docx)
    assert out["title"] == "My Title"


def test_extract_paragraph_styles(sample_docx: Path) -> None:
    out = extract(sample_docx)
    styles = {p["style"] for p in out["paragraphs"]}
    assert "Title" in styles
    assert "Heading 1" in styles
    assert "Normal" in styles


def test_extract_tables(sample_docx: Path) -> None:
    out = extract(sample_docx)
    assert len(out["tables"]) == 1
    assert out["tables"][0] == [["k1", "v1"], ["k2", "v2"]]


def test_extract_metadata(sample_docx: Path) -> None:
    out = extract(sample_docx)
    assert out["metadata"]["author"] == "Test Author"


def test_extract_normal_only_doc_captures_visual_signals(
    normal_only_docx: Path,
) -> None:
    """Every paragraph is style=Normal, but the title run is bold + 20pt. The new bold/max_font_pt fields must capture this so the LLM can spot the visual heading."""
    out = extract(normal_only_docx)
    assert all(p["style"] == "Normal" for p in out["paragraphs"])
    title_para = out["paragraphs"][0]
    assert title_para["bold"] is True
    assert title_para["max_font_pt"] == 20.0
    body_para = out["paragraphs"][1]
    assert body_para["bold"] is False
    assert body_para["max_font_pt"] is None


def test_extract_normal_only_doc_picks_visual_title(
    normal_only_docx: Path,
) -> None:
    """No Title/Heading 1 styles exist, so the heuristic must fall back to the bold paragraph with the largest font."""
    out = extract(normal_only_docx)
    assert out["title"] == "Standard Operating Procedure: Onboarding"
