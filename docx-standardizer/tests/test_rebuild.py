from pathlib import Path

import pytest
from docx import Document

from rebuild import rebuild
from schema import StandardizedDocument


@pytest.fixture
def fixture_master(tmp_path: Path) -> Path:
    p = tmp_path / "master.docx"
    doc = Document()
    doc.add_paragraph("master ref title", style="Title")
    doc.add_paragraph("master ref h1", style="Heading 1")
    doc.add_paragraph("master ref h2", style="Heading 2")
    doc.add_paragraph("master ref normal", style="Normal")
    doc.add_table(rows=1, cols=2, style="Table Grid")
    doc.save(str(p))
    return p


@pytest.fixture
def fixture_normalized() -> StandardizedDocument:
    return StandardizedDocument.model_validate(
        {
            "title": "Test Title",
            "document_control": {
                "doc_id": "DC-1",
                "owner": "Owner",
                "approval": "Approver",
                "effective_date": "2026-01-01",
            },
            "purpose": "A purpose.",
            "scope": "A scope.",
            "definitions": [],
            "responsibilities": [],
            "procedure": [{"step_number": 1, "action": "Step one."}],
            "records": [],
            "references": [],
            "revision_history": [
                {
                    "version": "1.0",
                    "date": "2026-01-01",
                    "author": "A",
                    "summary": "Init.",
                }
            ],
        }
    )


def test_rebuild_applies_styles(
    fixture_master: Path,
    fixture_normalized: StandardizedDocument,
    tmp_path: Path,
) -> None:
    doc = rebuild(fixture_normalized, fixture_master)
    out_path = tmp_path / "out.docx"
    doc.save(str(out_path))
    reloaded = Document(str(out_path))
    styles_used = {p.style.name for p in reloaded.paragraphs if p.text.strip()}
    assert "Title" in styles_used
    assert "Heading 1" in styles_used
    assert "Normal" in styles_used


def test_rebuild_includes_title(
    fixture_master: Path,
    fixture_normalized: StandardizedDocument,
    tmp_path: Path,
) -> None:
    doc = rebuild(fixture_normalized, fixture_master)
    out_path = tmp_path / "out.docx"
    doc.save(str(out_path))
    reloaded = Document(str(out_path))
    title_paras = [p for p in reloaded.paragraphs if p.style.name == "Title"]
    assert any(p.text == "Test Title" for p in title_paras)


def test_rebuild_writes_revision_history_table(
    fixture_master: Path,
    fixture_normalized: StandardizedDocument,
    tmp_path: Path,
) -> None:
    doc = rebuild(fixture_normalized, fixture_master)
    out_path = tmp_path / "out.docx"
    doc.save(str(out_path))
    reloaded = Document(str(out_path))
    revision_tables = [
        t for t in reloaded.tables if t.rows[0].cells[0].text == "Version"
    ]
    assert len(revision_tables) == 1
    assert revision_tables[0].rows[1].cells[3].text == "Init."
